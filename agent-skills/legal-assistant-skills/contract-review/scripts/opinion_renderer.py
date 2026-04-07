#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Render comprehensive review opinion text to a styled DOCX document.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DEFAULT_LINE_SPACING = 1.3
NUMBERED_ITEM_PATTERN = re.compile(r"(?:^|[\s：:；;])(\d+[\.、](?!\d))")


def render_opinion_docx(
    opinion_text: str,
    output_path: Path,
    font_name: str = "仿宋",
    base_font_size: int = 12,
    line_spacing: float = DEFAULT_LINE_SPACING,
    add_title: bool = True,
    title_text: str = "综合审核意见",
) -> Path:
    """
    Render opinion text into a DOCX file.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _set_document_font(doc, font_name, base_font_size, line_spacing)

    if add_title:
        _add_title(doc, title_text, font_name)

    paragraphs = _split_paragraphs(opinion_text)
    for para_text in paragraphs:
        for segment in _split_numbered_items(para_text):
            _add_paragraph(doc, segment, font_name, line_spacing)

    doc.save(output_path)
    return output_path


def _split_paragraphs(text: str) -> List[str]:
    raw = text.strip()
    if not raw:
        return []
    parts = [p.strip() for p in raw.split("\n\n") if p.strip()]
    if parts:
        return parts
    return [line.strip() for line in raw.splitlines() if line.strip()]


def _split_numbered_items(text: str) -> List[str]:
    cleaned = text.strip()
    if not cleaned:
        return []
    matches = list(NUMBERED_ITEM_PATTERN.finditer(cleaned))
    if not matches:
        return [cleaned]

    starts = [match.start(1) for match in matches]
    parts: List[str] = []
    lead = cleaned[:starts[0]].strip()
    if lead:
        parts.append(lead)
    for index, start in enumerate(starts):
        end = starts[index + 1] if index + 1 < len(starts) else len(cleaned)
        item = cleaned[start:end].strip()
        item = item.rstrip("；; ")
        if item:
            parts.append(item)
    return parts


def _set_document_font(doc: Document, font_name: str, base_font_size: int, line_spacing: float) -> None:
    for style_name in ("Normal", "Heading 1"):
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(base_font_size)
        _set_style_east_asia_font(style, font_name)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = line_spacing


def _set_style_east_asia_font(style, font_name: str) -> None:
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def _set_run_font(run, font_name: str, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def _add_title(doc: Document, text: str, font_name: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, font_name, size=16, bold=True)
    para.paragraph_format.space_after = Pt(6)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def _add_paragraph(doc: Document, text: str, font_name: str, line_spacing: float) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, font_name)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para.paragraph_format.line_spacing = line_spacing
