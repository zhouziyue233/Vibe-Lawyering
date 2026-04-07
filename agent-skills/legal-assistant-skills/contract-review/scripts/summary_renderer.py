#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Render contract summary text to a styled DOCX document.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


MAIN_SECTION_RE = re.compile(r"^[一二三四五六七八九十]+、")
SUB_SECTION_RE = re.compile(r"^\\d+\\.\\d+")


DEFAULT_LINE_SPACING = 1.3
DEFAULT_FIRST_COL_RATIO = 0.28
DEFAULT_CELL_MARGIN_TOP = 160
DEFAULT_CELL_MARGIN_BOTTOM = 120


def render_summary_docx(
    summary_text: str,
    output_path: Path,
    font_name: str = "仿宋",
    base_font_size: int = 12,
    line_spacing: float = DEFAULT_LINE_SPACING,
    first_col_ratio: float = DEFAULT_FIRST_COL_RATIO,
    cell_margin_top: int = DEFAULT_CELL_MARGIN_TOP,
    cell_margin_bottom: int = DEFAULT_CELL_MARGIN_BOTTOM,
) -> Path:
    """
    Render summary text into a DOCX file with headings and tables.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _set_document_font(doc, font_name, base_font_size, line_spacing)

    lines = [line.rstrip() for line in summary_text.strip().splitlines()]
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if _is_main_section(line):
            _add_heading(doc, line, level=1, font_name=font_name, line_spacing=line_spacing)
            i += 1
            continue

        if _is_sub_section(line):
            _add_heading(doc, line, level=2, font_name=font_name, line_spacing=line_spacing)
            i += 1
            continue

        if "\t" in line:
            rows, next_index = _collect_table_rows(lines, i)
            _add_table(
                doc,
                rows,
                font_name=font_name,
                line_spacing=line_spacing,
                first_col_ratio=first_col_ratio,
                cell_margin_top=cell_margin_top,
                cell_margin_bottom=cell_margin_bottom,
            )
            i = next_index
            continue

        if line.endswith("：") or line.endswith(":"):
            _add_label_paragraph(doc, line, font_name=font_name, line_spacing=line_spacing)
            i += 1
            continue

        _add_paragraph(doc, line, font_name=font_name, line_spacing=line_spacing)
        i += 1

    _add_page_numbers(doc, font_name, base_font_size)
    doc.save(output_path)
    return output_path


def _set_document_font(doc: Document, font_name: str, base_font_size: int, line_spacing: float) -> None:
    for style_name in ("Normal", "Heading 1", "Heading 2"):
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(base_font_size)
        _set_style_east_asia_font(style, font_name)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = line_spacing


def _set_style_east_asia_font(style, font_name: str) -> None:
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def _set_run_font(run, font_name: str, size: Optional[int] = None, bold: Optional[bool] = None) -> None:
    run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def _add_heading(doc: Document, text: str, level: int, font_name: str, line_spacing: float) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    size = 16 if level == 1 else 14
    _set_run_font(run, font_name, size=size, bold=True)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(4)
    _apply_line_spacing(para, line_spacing)
    if level == 1:
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def _add_paragraph(doc: Document, text: str, font_name: str, line_spacing: float) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, font_name)
    para.paragraph_format.space_after = Pt(2)
    _apply_line_spacing(para, line_spacing)


def _add_label_paragraph(doc: Document, text: str, font_name: str, line_spacing: float) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, font_name, bold=True)
    para.paragraph_format.space_after = Pt(2)
    _apply_line_spacing(para, line_spacing)


def _add_table(
    doc: Document,
    rows: List[Tuple[str, str]],
    font_name: str,
    line_spacing: float,
    first_col_ratio: float,
    cell_margin_top: int,
    cell_margin_bottom: int,
) -> None:
    if not rows:
        return
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    _set_table_column_widths(doc, table, first_col_ratio)
    for idx, (left, right) in enumerate(rows):
        cells = table.add_row().cells
        _set_cell_margins(cells[0], top=cell_margin_top, bottom=cell_margin_bottom)
        _set_cell_margins(cells[1], top=cell_margin_top, bottom=cell_margin_bottom)
        _set_cell_text(cells[0], left, font_name=font_name, bold=(idx == 0), line_spacing=line_spacing)
        _set_cell_text(cells[1], right, font_name=font_name, bold=(idx == 0), line_spacing=line_spacing)


def _set_cell_text(cell, text: str, font_name: str, bold: bool = False, line_spacing: float = DEFAULT_LINE_SPACING) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    _set_run_font(run, font_name, bold=bold)
    _apply_line_spacing(para, line_spacing)


def _apply_line_spacing(para, line_spacing: float) -> None:
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para.paragraph_format.line_spacing = line_spacing


def _set_table_column_widths(doc: Document, table, first_col_ratio: float) -> None:
    if not table.columns or len(table.columns) < 2:
        return
    section = doc.sections[0]
    content_width = section.page_width - section.left_margin - section.right_margin
    first_width = int(content_width * first_col_ratio)
    second_width = int(content_width - first_width)
    table.columns[0].width = first_width
    table.columns[1].width = second_width
    for row in table.rows:
        row.cells[0].width = first_width
        row.cells[1].width = second_width


def _set_cell_margins(cell, top: int | None = None, bottom: int | None = None, left: int | None = None, right: int | None = None) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    _set_tc_margin(tc_mar, "w:top", top)
    _set_tc_margin(tc_mar, "w:bottom", bottom)
    _set_tc_margin(tc_mar, "w:left", left)
    _set_tc_margin(tc_mar, "w:right", right)


def _set_tc_margin(tc_mar, tag: str, value: int | None) -> None:
    if value is None:
        return
    node = tc_mar.find(qn(tag))
    if node is None:
        node = OxmlElement(tag)
        tc_mar.append(node)
    node.set(qn("w:w"), str(value))
    node.set(qn("w:type"), "dxa")


def _add_page_numbers(doc: Document, font_name: str, base_font_size: int) -> None:
    for section in doc.sections:
        footer = section.footer
        if footer.paragraphs:
            para = footer.paragraphs[0]
        else:
            para = footer.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        run = para.add_run()
        _set_run_font(run, font_name, size=base_font_size)

        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_char_begin)

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = "PAGE"
        run._r.append(instr_text)

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char_end)


def _collect_table_rows(lines: List[str], start_index: int) -> Tuple[List[Tuple[str, str]], int]:
    rows: List[Tuple[str, str]] = []
    i = start_index
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            break
        if _is_main_section(line) or _is_sub_section(line):
            break
        if "\t" not in line:
            break
        left, right = line.split("\t", 1)
        rows.append((left.strip(), right.strip()))
        i += 1
    return rows, i


def _is_main_section(line: str) -> bool:
    return bool(MAIN_SECTION_RE.match(line))


def _is_sub_section(line: str) -> bool:
    return bool(SUB_SECTION_RE.match(line))
